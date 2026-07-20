"""Backend API tests for WANOSC-Toolbox (auth, admin users, Word→Excel, Word→PDF)."""
import io
import os
import uuid
import pytest
import requests
from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader

BASE_URL = os.environ.get("REACT_APP_BACKEND_URL", "").rstrip("/")
if not BASE_URL:
    with open("/app/frontend/.env") as f:
        for line in f:
            if line.startswith("REACT_APP_BACKEND_URL="):
                BASE_URL = line.split("=", 1)[1].strip().rstrip("/")
API = f"{BASE_URL}/api"


def make_docx(title="TEST_Doc Title", body="Hello world.", h1="Intro"):
    doc = Document()
    doc.core_properties.title = title
    doc.add_paragraph(title, style="Title")
    doc.add_paragraph(h1, style="Heading 1")
    doc.add_paragraph(body)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# --- fixtures ---
@pytest.fixture(scope="module")
def admin_token():
    r = requests.post(f"{API}/auth/login", json={"username": "admin", "password": "admin123"})
    assert r.status_code == 200, r.text
    data = r.json()
    assert data["user"]["role"] == "admin"
    return data["token"]


@pytest.fixture(scope="module")
def admin_headers(admin_token):
    return {"Authorization": f"Bearer {admin_token}"}


# --- Auth ---
class TestAuth:
    def test_login_valid_admin(self):
        r = requests.post(f"{API}/auth/login", json={"username": "admin", "password": "admin123"})
        assert r.status_code == 200
        d = r.json()
        assert "token" in d and d["user"]["username"] == "admin" and d["user"]["role"] == "admin"

    def test_login_invalid(self):
        r = requests.post(f"{API}/auth/login", json={"username": "admin", "password": "wrong"})
        assert r.status_code == 401

    def test_me_unauthenticated(self):
        r = requests.get(f"{API}/auth/me")
        assert r.status_code == 401

    def test_me_authenticated(self, admin_headers):
        r = requests.get(f"{API}/auth/me", headers=admin_headers)
        assert r.status_code == 200
        assert r.json()["role"] == "admin"


# --- Admin user management ---
class TestAdminUsers:
    _test_username = f"TEST_user_{uuid.uuid4().hex[:8]}"
    _test_user_id = None
    _general_token = None

    def test_list_users_admin(self, admin_headers):
        r = requests.get(f"{API}/users", headers=admin_headers)
        assert r.status_code == 200
        users = r.json()
        assert isinstance(users, list)
        assert any(u["username"] == "admin" for u in users)

    def test_create_general_user(self, admin_headers):
        payload = {"username": self._test_username, "password": "pass1234", "name": "Test User", "role": "user"}
        r = requests.post(f"{API}/users", json=payload, headers=admin_headers)
        assert r.status_code == 200, r.text
        u = r.json()
        assert u["username"] == self._test_username.lower()
        assert u["role"] == "user"
        TestAdminUsers._test_user_id = u["id"]

    def test_new_user_can_login(self):
        r = requests.post(f"{API}/auth/login", json={"username": self._test_username, "password": "pass1234"})
        assert r.status_code == 200
        d = r.json()
        assert d["user"]["role"] == "user"
        TestAdminUsers._general_token = d["token"]

    def test_general_user_forbidden_from_users_endpoints(self):
        assert TestAdminUsers._general_token, "need general token"
        h = {"Authorization": f"Bearer {TestAdminUsers._general_token}"}
        assert requests.get(f"{API}/users", headers=h).status_code == 403
        assert requests.post(f"{API}/users", headers=h,
                             json={"username": "x", "password": "y", "role": "user"}).status_code == 403
        assert requests.delete(f"{API}/users/999999", headers=h).status_code == 403

    def test_unauthenticated_forbidden(self):
        assert requests.get(f"{API}/users").status_code == 401
        assert requests.post(f"{API}/users", json={"username": "x", "password": "y"}).status_code == 401

    def test_create_duplicate_username_400(self, admin_headers):
        payload = {"username": self._test_username, "password": "x", "role": "user"}
        r = requests.post(f"{API}/users", json=payload, headers=admin_headers)
        assert r.status_code == 400

    def test_admin_cannot_delete_self(self, admin_headers):
        me = requests.get(f"{API}/auth/me", headers=admin_headers).json()
        r = requests.delete(f"{API}/users/{me['id']}", headers=admin_headers)
        assert r.status_code == 400

    def test_delete_user(self, admin_headers):
        assert TestAdminUsers._test_user_id
        r = requests.delete(f"{API}/users/{TestAdminUsers._test_user_id}", headers=admin_headers)
        assert r.status_code == 200
        # Verify gone
        users = requests.get(f"{API}/users", headers=admin_headers).json()
        assert not any(u["id"] == TestAdminUsers._test_user_id for u in users)


# --- Word -> Excel /api/convert ---
class TestConvert:
    def test_convert_requires_auth(self):
        r = requests.post(f"{API}/convert",
                          files=[("files", ("a.docx", make_docx(), "application/octet-stream"))])
        assert r.status_code == 401

    def test_convert_single_docx(self, admin_headers):
        content = make_docx(title="TEST_T1", body="Body text")
        files = [("files", ("t1.docx", content, "application/octet-stream"))]
        r = requests.post(f"{API}/convert", files=files, headers=admin_headers)
        assert r.status_code == 200, r.text
        data = r.json()
        assert data["file_count"] == 1
        assert "t1.docx" in data["filenames"]
        assert "job_id" in data
        # FACT_COLUMNS schema
        assert "Fact#" in data["columns"] and "Facts" in data["columns"]
        pytest.job_id = data["job_id"]

    def test_convert_mixed_valid_and_invalid(self, admin_headers):
        files = [
            ("files", ("good.docx", make_docx(body="Good"), "application/octet-stream")),
            ("files", ("bad.txt", b"not a docx", "text/plain")),
        ]
        r = requests.post(f"{API}/convert", files=files, headers=admin_headers)
        assert r.status_code == 200
        data = r.json()
        assert data["file_count"] == 1
        assert any(e["filename"] == "bad.txt" for e in data["errors"])

    def test_convert_all_invalid_returns_400(self, admin_headers):
        files = [("files", ("bad.txt", b"nope", "text/plain"))]
        r = requests.post(f"{API}/convert", files=files, headers=admin_headers)
        assert r.status_code == 400


class TestJobs:
    @pytest.fixture(scope="class")
    def job_id(self, admin_headers):
        r = requests.post(f"{API}/convert",
                          files=[("files", ("fx.docx", make_docx(body="fx"), "application/octet-stream"))],
                          headers=admin_headers)
        assert r.status_code == 200
        return r.json()["job_id"]

    def test_list_jobs(self, admin_headers, job_id):
        r = requests.get(f"{API}/jobs", headers=admin_headers)
        assert r.status_code == 200
        assert any(j["job_id"] == job_id for j in r.json())

    def test_get_job(self, admin_headers, job_id):
        r = requests.get(f"{API}/jobs/{job_id}", headers=admin_headers)
        assert r.status_code == 200
        assert r.json()["job_id"] == job_id

    def test_download_excel(self, admin_headers, job_id):
        r = requests.get(f"{API}/jobs/{job_id}/download", headers=admin_headers)
        assert r.status_code == 200
        assert "spreadsheetml" in r.headers.get("Content-Type", "")
        wb = load_workbook(io.BytesIO(r.content))
        ws = wb.active
        headers_row = [c.value for c in ws[1]]
        for expected in ["Fact#", "Facts", "Title", "Scope", "File Name", "Keyword", "Area", "Team Area", "Reviewer"]:
            assert expected in headers_row, f"missing {expected}"

    def test_delete_job(self, admin_headers):
        r = requests.post(f"{API}/convert",
                          files=[("files", ("del.docx", make_docx(body="del"), "application/octet-stream"))],
                          headers=admin_headers)
        jid = r.json()["job_id"]
        assert requests.delete(f"{API}/jobs/{jid}", headers=admin_headers).status_code == 200
        assert requests.get(f"{API}/jobs/{jid}", headers=admin_headers).status_code == 404


# --- Word -> PDF ---
class TestWordToPdf:
    def test_pdf_requires_auth(self):
        r = requests.post(f"{API}/word-to-pdf",
                          files=[("files", ("a.docx", make_docx(), "application/octet-stream"))])
        assert r.status_code == 401

    def test_pdf_combine_two_docx(self, admin_headers):
        files = [
            ("files", ("one.docx", make_docx(body="Body one"), "application/octet-stream")),
            ("files", ("two.docx", make_docx(body="Body two"), "application/octet-stream")),
        ]
        r = requests.post(f"{API}/word-to-pdf", files=files, headers=admin_headers, timeout=120)
        assert r.status_code == 200, r.text
        assert r.headers.get("Content-Type", "").startswith("application/pdf")
        assert r.headers.get("X-Converted-Count") == "2"
        assert r.headers.get("X-Error-Count") == "0"
        # Verify PDF parseable
        reader = PdfReader(io.BytesIO(r.content))
        assert len(reader.pages) >= 2

    def test_pdf_skips_non_docx(self, admin_headers):
        files = [
            ("files", ("good.docx", make_docx(body="Good"), "application/octet-stream")),
            ("files", ("bad.txt", b"not a docx", "text/plain")),
        ]
        r = requests.post(f"{API}/word-to-pdf", files=files, headers=admin_headers, timeout=120)
        assert r.status_code == 200
        assert r.headers.get("X-Converted-Count") == "1"
        assert r.headers.get("X-Error-Count") == "1"

    def test_pdf_all_invalid_returns_400(self, admin_headers):
        files = [("files", ("bad.txt", b"nope", "text/plain"))]
        r = requests.post(f"{API}/word-to-pdf", files=files, headers=admin_headers)
        assert r.status_code == 400
